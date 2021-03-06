#!/usr/bin/env python
# -*-coding:utf-8-*-
"""
author : shenshuo
date   : 2017年11月20日11:10:29
role   : 用户管理API

status = '0'    正常
status = '10'   逻辑删除
status = '20'   禁用
"""

import json
import shortuuid
import base64
from websdk.jwt_token import gen_md5
from websdk.tools import check_password
from libs.base_handler import BaseHandler
from websdk.db_context import DBContext
# from models.admin import Users, UserRoles,model_to_dict as users_model_to_dict
from models.problem import CaseList, CaseUsers, model_to_dict
from models.problem import PlanList, model_to_dict  as   model_to_dict2
from websdk.consts import const
from websdk.cache_context import cache_conn
from websdk.web_logs import ins_log
from sqlalchemy import or_, and_
from sqlalchemy import func
import tornado.ioloop
import tornado.web
import pandas as pd
import matplotlib.pyplot as plt
import numpy as np
from docx import shared
import os
import time,datetime

class CaseListHandler(BaseHandler):
    def post(self, *args, **kwargs):
        data = json.loads(self.request.body.decode("utf-8"))
        case_name = data.get('case_name', None)  # 个案名称
        case_priority = data.get('case_priority', None)  # 优先级
        case_executor = data.get('case_executor', False)  # 执行人
        case_type = data.get('case_type', False)  # 类型
        demander = data.get('demander', False) # 需求人
        if "--" in demander:
            demander =  demander.split("--")[0]
        case_details = data.get('case_details', False)  # 详情描述
        case_source = data.get('case_source', False)  # 来源

        demand_unit = data.get('demand_unit', False)  # 需求单位
        case_status = data.get('case_status', False)  # 状态
        case_obj = data.get('case_obj', False)  # 项目
        case_stime = data.get('case_stime', False)
        case_etime = data.get('case_etime', False)
        case_creator = data.get('case_creator', False)
        #case_ltime = data.get('case_ltime', False)
        stimeArray = time.strptime(case_stime,"%Y-%m-%d %H:%M:%S")
        stimeStamp = int(time.mktime(stimeArray))
        etimeArray = time.strptime(case_etime, "%Y-%m-%d %H:%M:%S")
        etimeStamp = int(time.mktime(etimeArray))
        case_ltime = int((etimeStamp - stimeStamp)/60)

        import uuid
        temp = str(uuid.uuid1().int >> 64)
        with DBContext('w', None, True) as session:
            session.add(CaseList(
                case_num=temp,
                case_name=case_name,
                case_priority=case_priority,
                case_executor=case_executor,
                case_type=case_type,
                demander=demander,
                case_details=case_details,
                case_source=case_source,
                demand_unit=demand_unit,
                case_status=case_status,
                case_obj=case_obj,
                case_stime=case_stime,
                case_creator=case_creator,
                case_etime=case_etime,
                case_ltime=case_ltime))
            session.commit()
        self.write(dict(code=0, msg='成功', count=0, data=[]))

    def put(self, *args, **kwargs):
        data = json.loads(self.request.body.decode("utf-8"))
        case_num = data.get('case_num', None)  # 编号
        case_id = int(data.get('case_id', None))  # id号
        case_name = data.get('case_name', None)  # 个案名称
        case_priority = data.get('case_priority', None)  # 优先级
        case_executor = data.get('case_executor', False)  # 执行人
        case_type = data.get('case_type', False)  # 类型
        demander = str(data.get('demander', False))  # 需求人
        if "--" in demander:
            demander =  demander.split("--")[0]
        case_details = data.get('case_details', False)  # 详情描述
        case_source = data.get('case_source', False)  # 来源
        demand_unit = data.get('demand_unit', False)  # 需求单位
        case_status = data.get('case_status', False)  # 状态
        case_obj = data.get('case_obj', False)  # 项目
        case_stime = data.get('case_stime', False)
        case_etime = data.get('case_etime', False)
        case_creator = data.get('case_creator', False)
        # case_ltime = data.get('case_ltime', False)  #
        stimeArray = time.strptime(case_stime, "%Y-%m-%d %H:%M:%S")
        stimeStamp = int(time.mktime(stimeArray))
        etimeArray = time.strptime(case_etime, "%Y-%m-%d %H:%M:%S")
        etimeStamp = int(time.mktime(etimeArray))
        case_ltime = int((etimeStamp - stimeStamp) / 60)

        with DBContext('w', None, True) as session:
            session.query(CaseList).filter(CaseList.id == case_id).update({
                CaseList.case_name: case_name,
                CaseList.case_priority: case_priority,
                CaseList.case_executor: case_executor,
                CaseList.case_type: case_type,
                CaseList.demander: demander,
                CaseList.case_details: case_details,
                CaseList.case_source: case_source,
                CaseList.demand_unit: demand_unit,
                CaseList.case_status: case_status,
                CaseList.case_obj: case_obj,
                CaseList.case_stime: case_stime,
                CaseList.case_etime: case_etime,
                CaseList.case_creator: case_creator,
                CaseList.case_ltime: case_ltime,
            })

        self.write(dict(code=0, msg='成功', count=0, data=[]))


class getCaseListHandler(BaseHandler):
    def get(self, *args, **kwargs):
        data_list = []
        superuser_flag = 0
        # username = self.get_current_user()
        nickname = self.get_current_nickname()
        # toname = self.get_argument('key', strip=True)  # 要查询的字段
        # tovalue = self.get_argument('value', strip=True)
        tovalue = self.get_argument('value', strip=True)  # 要查询的关键字
        topage = int(self.get_argument('page', strip=1))  # 开始页
        tolimit = int(self.get_argument('limit', strip=10))  # 要查询条数
        isExport = self.get_argument('isExport')  # 是否导出
        limit_start = (topage - 1) * tolimit

        with DBContext('r') as session:
            params = {}
            if tovalue:
                params = eval(tovalue)
            conditions = []
            if self.is_superuser:
                superuser_flag = 1
                if params.get('case_executor', ''):
                    conditions.append(CaseList.case_executor  == params['case_executor'])
            else:
                conditions.append(or_(CaseList.case_executor == nickname, CaseList.case_creator == nickname))
            if params.get('case_name', ''):
                conditions.append(CaseList.case_name.like('%{}%'.format(params['case_name'])))
            if params.get('case_details', ''):
                conditions.append(CaseList.case_details.like('%{}%'.format(params['case_details'])))
            if params.get('case_type', ''):
                conditions.append(CaseList.case_type == params['case_type'])
            if params.get('case_status', ''):
                conditions.append(CaseList.case_status == params['case_status'])
            if params.get('case_priority', ''):
                conditions.append(CaseList.case_priority == params['case_priority'])
            if params.get('demander', ''):
                conditions.append(CaseList.demander.like('%{}%'.format(params['demander'])))
            if params.get('case_source', ''):
                conditions.append(CaseList.case_source == params['case_source'])
            if params.get('case_obj', ''):
                conditions.append(CaseList.case_obj == params['case_obj'])
            if params.get('demand_unit', ''):
                conditions.append(CaseList.demand_unit == params['demand_unit'])
            # if params.get('case_stime', ''):
            #     temptimestr = params['case_stime'] + " 00:00:00"
            #     conditions.append(CaseList.ctime >= temptimestr)
            # if params.get('case_etime', ''):
            #     temptimestr = params['case_etime'] + "  23:59:59"
            #     conditions.append(CaseList.ctime <= temptimestr)
            if params.get('case_stime', '') and params.get('case_etime', ''):
                stemptimestr = params['case_stime'] + " 00:00:00"
                etemptimestr = params['case_etime'] + "  23:59:59"
                conditions.append(or_(and_(CaseList.case_stime >=  stemptimestr, CaseList.case_etime <= etemptimestr),
                                      and_(CaseList.case_stime < stemptimestr, CaseList.case_etime >= stemptimestr),
                                      and_(CaseList.case_stime <= etemptimestr, CaseList.case_etime > etemptimestr),
                                      and_(CaseList.case_stime < stemptimestr, CaseList.case_etime > etemptimestr),
                                      ))

            if isExport != 'false':
                todata = session.query(CaseList).filter(*conditions).order_by(CaseList.ctime.desc()).all()
                tocount = session.query(CaseList).filter(*conditions).count()
            else:
                todata = session.query(CaseList).filter(*conditions).order_by(CaseList.ctime.desc()).offset(
                    limit_start).limit(int(tolimit)).all()
                tocount = session.query(CaseList).filter(*conditions).count()

        for msg in todata:
            case_dict = {}
            data_dict = model_to_dict(msg)
            case_dict["id"] = data_dict["id"]
            case_dict["case_num"] = data_dict["case_num"]
            case_dict["case_obj"] = data_dict["case_obj"]
            case_dict["demand_unit"] = data_dict["demand_unit"]
            case_dict["case_details"] = data_dict["case_details"]
            case_dict["case_type"] = data_dict["case_type"]
            case_dict["case_ltime"] = data_dict["case_ltime"]
            case_dict["case_name"] = data_dict["case_name"]
            case_dict["case_status"] = data_dict["case_status"]
            case_dict["case_priority"] = data_dict["case_priority"]
            case_dict["demander"] = data_dict["demander"]
            case_dict["case_executor"] = data_dict["case_executor"]
            case_dict["case_source"] = data_dict["case_source"]
            case_dict["case_details"] = data_dict["case_details"]
            case_dict["case_stime"] = str(data_dict["case_stime"])
            case_dict["case_etime"] = str(data_dict["case_etime"])
            case_dict["case_creator"] = data_dict["case_creator"]
            data_list.append(case_dict)

        if len(data_list) > 0:
            self.write(dict(code=0, msg='获取成功', count=tocount, data=data_list ,flag = superuser_flag ))
        else:
            self.write(dict(code=-1, msg='没有相关数据', count=0, data=[],flag = superuser_flag))


class getCasefileHandler(BaseHandler):
    def get(self, *args, **kwargs):
        data_list = []
        nickname = self.get_current_nickname()
        tostart = self.get_argument('startdate', strip=True)  # 要查询的关键字
        toend = self.get_argument('enddate', strip=True)  # 要查询的关键字
        # CaseList表
        with DBContext('r') as session:
            params = {}
            conditions = []
            # if self.is_superuser:
            #     pass
            # else:
            #     conditions.append(CaseList.case_executor == nickname)
            conditions.append(CaseList.case_stime >= tostart)
            conditions.append(CaseList.case_etime <= toend)
            todata = session.query(CaseList).filter(*conditions).order_by(CaseList.ctime.desc()).all()
        for msg in todata:
            case_dict = {}
            data_dict = model_to_dict(msg)
            case_dict["id"] = data_dict["id"]
            case_dict["case_num"] = data_dict["case_num"]
            case_dict["case_obj"] = data_dict["case_obj"]
            case_dict["demand_unit"] = data_dict["demand_unit"]
            case_dict["case_details"] = data_dict["case_details"]
            case_dict["case_type"] = data_dict["case_type"]
            case_dict["case_ltime"] = data_dict["case_ltime"]
            case_dict["case_name"] = data_dict["case_name"]
            case_dict["case_status"] = data_dict["case_status"]
            case_dict["case_priority"] = data_dict["case_priority"]
            case_dict["demander"] = data_dict["demander"]
            case_dict["case_executor"] = data_dict["case_executor"]
            case_dict["case_source"] = data_dict["case_source"]
            case_dict["case_details"] = data_dict["case_details"]
            case_dict["case_stime"] = str(data_dict["case_stime"])
            case_dict["case_etime"] = str(data_dict["case_etime"])
            case_dict["case_creator"] = data_dict["case_creator"]
            data_list.append(case_dict)
        # PlanList表
        plandata_list = []  # 计划工作完成情况
        with DBContext('r') as session:
            params = {}
            conditions = []
            # if self.is_superuser:
            #     pass
            # else:
            #     conditions.append(CaseList.case_executor == nickname)
            conditions.append(PlanList.plan_stime >= tostart)
            conditions.append(PlanList.plan_etime <= toend)
            conditions.append(PlanList.plan_status == "处理中")
            todata = session.query(PlanList).filter(*conditions).order_by(PlanList.ctime.desc()).all()
        for msg in todata:
            plan_dict = {}
            data_dict = model_to_dict(msg)
            plan_dict["id"] = data_dict["id"]
            plan_dict["plan_num"] = data_dict["plan_num"]
            plan_dict["plan_obj"] = data_dict["plan_obj"]
            plan_dict["demand_unit"] = data_dict["demand_unit"]
            plan_dict["plan_details"] = data_dict["plan_details"]
            plan_dict["plan_type"] = data_dict["plan_type"]
            plan_dict["plan_ltime"] = data_dict["plan_ltime"]
            plan_dict["plan_name"] = data_dict["plan_name"]
            plan_dict["plan_status"] = data_dict["plan_status"]
            plan_dict["plan_priority"] = data_dict["plan_priority"]
            plan_dict["demander"] = data_dict["demander"]
            plan_dict["plan_executor"] = data_dict["plan_executor"]
            plan_dict["plan_source"] = data_dict["plan_source"]
            plan_dict["plan_details"] = data_dict["plan_details"]
            plan_dict["plan_stime"] = str(data_dict["plan_stime"])
            plan_dict["plan_etime"] = str(data_dict["plan_etime"])
            plan_dict["plan_creator"] = data_dict["plan_creator"]
            plandata_list.append(plan_dict)
        ins_log.read_log('info', plandata_list)
        if (len(data_list) + len(plandata_list)) > 0:
            import docx
            import time, datetime
            flag = 0  # 周报0月报1
            # ins_log.read_log('info', "800000000000000000000000000000000000")
            Base_DIR = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
            upload_path = '{}/static/report/template/运维组工作报告模板.docx'.format(Base_DIR)
            # ins_log.read_log('info', upload_path)
            # ins_log.read_log('info', "8000000000011111111111111111111")
            # doc = docx.Document(u"/opt/codo/codo-problem/static/report/template/运维组工作报告模板.docx")
            doc = docx.Document(upload_path)
            if (int(str(toend)[8:10]) - int(str(tostart)[8:10])) <= 7:
                doc.paragraphs[0].text = ['工作周报']
                flag = 0
            else:
                doc.paragraphs[0].text = ['工作月报']
                flag = 1
            tempstr = '报告时间：' + str(tostart)[0:4] + '年' + str(tostart)[5:7] + '月' + str(tostart)[8:10] + '日' \
                      + '～' + str(toend)[0:4] + '年' + str(toend)[5:7] + '月' + str(toend)[8:10] + '日'
            # ins_log.read_log('info', tempstr)
            tempdict = []
            tempdict.append(tempstr)
            # ins_log.read_log('info', tempdict)
            doc.paragraphs[3].text = tempdict
            tempstr = ''
            tempstr = '报 告 人：' + nickname
            tempdict = []
            tempdict.append(tempstr)
            doc.paragraphs[4].text = tempdict
            # 本周/月工作完成情况（2020-05-21~2020-05-31）
            tempstr = ''
            if flag == 0:
                tempstr = '本周工作完成情况（'
            else:
                tempstr = '本月工作完成情况（'

            tempstr = tempstr + str(tostart)[0:4] + '-' + str(tostart)[5:7] + '-' + str(tostart)[8:10] \
                      + '～' + str(toend)[0:4] + '-' + str(toend)[5:7] + '-' + str(toend)[8:10] + '）'
            tempdict = []
            tempdict.append(tempstr)
            doc.paragraphs[6].text = tempdict
            # 下周/月工作计划（2020-xx-xx~2020-xx-xx）
            # ins_log.read_log('info', toend)
            if flag == 0:
                tempstr = ''
                tempstr = toend + ' ' + '23:59:59'
                # ins_log.read_log('info', tempstr)
                timeArray = time.strptime(tempstr, "%Y-%m-%d %H:%M:%S")
                # 转为时间戳
                timeStamp = int(time.mktime(timeArray))
                # ins_log.read_log('info', timeStamp)
                # 下周一
                timeStamp1 = timeStamp + 3600 * 24
                startlocaltime = time.localtime(timeStamp1)
                startdatatime = time.strftime("%Y-%m-%d %H:%M:%S", startlocaltime)
                # ins_log.read_log('info', startdatatime)
                # 下周日
                timeStamp2 = timeStamp + 3600 * 24 * 7
                endlocaltime = time.localtime(timeStamp2)
                enddatatime = time.strftime("%Y-%m-%d %H:%M:%S", endlocaltime)
                # ins_log.read_log('info', enddatatime)
                tempstr = ''
                tempstr = '下周工作计划（' + str(startdatatime)[0:4] + '-' + str(startdatatime)[5:7] + '-' + str(
                    startdatatime)[8:10] \
                          + '～' + str(enddatatime)[0:4] + '-' + str(enddatatime)[5:7] + '-' + str(enddatatime)[
                                                                                              8:10] + '）'
            else:
                import calendar
                toyear = int(str(toend)[0:4])
                tomonth = int(str(toend)[5:7]) + 1
                if tomonth > 12:
                    tomonth = 1
                    toyear = toyear + 1
                monthRange = calendar.monthrange(toyear, tomonth)  # 返回值是元组，第一个参数是这个月的第一天是星期几，第二个参数是这个月的总天数

                tempstr = ''
                if len(str(tomonth)) == 1:
                    tomonthstr = '0' + str(tomonth)
                else:
                    tomonthstr = str(tomonth)
                tempstr = '下个月工作计划（' + str(toyear) + '-' + tomonthstr + '-' + "01" \
                          + '～' + str(toyear) + '-' + tomonthstr + '-' + str(monthRange[1]) + '）'
            tempdict = []
            tempdict.append(tempstr)
            doc.paragraphs[7].text = tempdict

            rows_index = 0  # 行数
            merge_index = 0  # 合并个数
            totable = doc.tables[0]
            laiwen_list = []  # 来问
            upgrade_list = []  # 升级
            fault_list = []  # 故障

            especially_list = []  # 特急/耗时长的情况
            sudden_list = []  # 其他突发工作
            # data_list.extend(plan_list) #合并两个列表
            for i in data_list:
                if i["case_source"] == '来文':
                    laiwen_list.append(i)
                    totable.add_row()
                if i["case_type"] == '应用升级':
                    upgrade_list.append(i)
                    totable.add_row()
                if i["case_type"] == '故障':
                    fault_list.append(i)
                    totable.add_row()
                if i["case_priority"] == '特急' or int(i["case_ltime"]) >= 240:
                    especially_list.append(i)
                    totable.add_row()
            for i in plandata_list:
                totable.add_row()

            rows_index += 1  # 行数
            merge_index += len(laiwen_list)  # 合并个数
            if len(laiwen_list) == 0:
                totable.add_row()
                totable.cell(rows_index, 1).text = "来文情况"
                rows_index += 1  # 行数
                merge_index += 1  # 合并个数
            else:
                for k in range(0, len(laiwen_list)):
                    totable.cell(rows_index, 1).text = "来文情况"
                    totable.cell(k + 1, 2).text = str(k + 1) + '.' + laiwen_list[k]["case_name"]
                totable.cell(rows_index, 1).merge(totable.cell(len(laiwen_list), 1))

            rows_index += len(laiwen_list)  # 行数
            merge_index += len(upgrade_list)  # 合并个数
            if len(upgrade_list) == 0:
                totable.add_row()
                totable.cell(rows_index, 1).text = "升级情况"
                rows_index += 1  # 行数
                merge_index += 1  # 合并个数
            else:
                for j in range(0, len(upgrade_list)):
                    totable.cell(rows_index, 1).text = "升级情况"
                    totable.cell(rows_index + j, 2).text = str(j + 1) + '.' + upgrade_list[j]["case_name"]
                totable.cell(rows_index, 1).merge(totable.cell(merge_index, 1))

            rows_index += len(upgrade_list)  # 行数
            merge_index += len(fault_list)  # 合并个数
            if len(fault_list) == 0:
                totable.add_row()
                totable.cell(rows_index, 1).text = "故障情况"
                rows_index += 1  # 行数
                merge_index += 1  # 合并个数
            else:
                for h in range(0, len(fault_list)):
                    totable.cell(rows_index, 1).text = "故障情况"
                    totable.cell(rows_index + h, 2).text = str(h + 1) + '.' + fault_list[h]["case_name"]
                totable.cell(rows_index, 1).merge(totable.cell(merge_index, 1))

            rows_index += len(fault_list)  # 行数
            merge_index += len(plandata_list)  # 合并个数
            if len(plandata_list) == 0:
                totable.add_row()
                totable.cell(rows_index, 1).text = "计划工作完成情况"
                rows_index += 1  # 行数
                merge_index += 1  # 合并个数
            else:
                for g in range(0, len(plandata_list)):
                    totable.cell(rows_index, 1).text = "计划工作完成情况"
                    totable.cell(rows_index + g, 2).text = str(g + 1) + '.' + plandata_list[g]["plan_name"]
                totable.cell(rows_index, 1).merge(totable.cell(merge_index, 1))

            rows_index += len(plandata_list)  # 行数
            merge_index += len(especially_list)  # 合并个数
            if len(especially_list) == 0:
                totable.add_row()
                totable.cell(rows_index, 1).text = "重要工作情况"
                rows_index += 1  # 行数
                merge_index += 1  # 合并个数
            else:
                for g in range(0, len(especially_list)):
                    totable.cell(rows_index, 1).text = "重要工作情况"
                    totable.cell(rows_index + g, 2).text = str(g + 1) + '.' + especially_list[g]["case_name"]
                totable.cell(rows_index, 1).merge(totable.cell(merge_index, 1))

            rows_index += len(especially_list)  # 行数
            merge_index += len(sudden_list)  # 合并个数
            if len(sudden_list) == 0:
                totable.add_row()
                totable.cell(rows_index, 1).text = "其他突发工作"
                rows_index += 1  # 行数
                merge_index += 1  # 合并个数
            else:
                for g in range(0, len(sudden_list)):
                    totable.cell(rows_index, 1).text = "其他突发工作"
                    totable.cell(rows_index + g, 2).text = str(g + 1) + '.' + sudden_list[g]["case_name"]
                totable.cell(rows_index, 1).merge(totable.cell(merge_index, 1))

            # ins_log.read_log('info', especially_list)
            # ins_log.read_log('info', len(especially_list))
            totable.cell(1, 0).merge(totable.cell(merge_index, 0))

            tempstr = ''
            tempstr = '维护组工作报告_' + nickname + '[' + str(tostart)[0:4] + str(tostart)[5:7] + str(tostart)[8:10] \
                      + '-' + str(toend)[0:4] + str(toend)[5:7] + str(toend)[8:10] + ']' + '.docx'
            # ins_log.read_log('info', tempstr)
            # 个案数据统计分析
            with DBContext('r') as session:
                conditions = []
                conditions.append(CaseList.case_source == "来文")
                conditions.append(CaseList.case_stime >= tostart)
                conditions.append(CaseList.case_etime <= toend)
                lime = session.query(func.count(CaseList.case_source),
                                     func.date_format(CaseList.case_stime, '%Y-%m-%d').label('date')).filter(
                    *conditions).group_by('date').all()
            with DBContext('r') as session:
                conditions = []
                conditions.append(CaseList.case_type == "应用升级")
                conditions.append(CaseList.case_stime >= tostart)
                conditions.append(CaseList.case_etime <= toend)
                lime2 = session.query(func.count(CaseList.case_type),
                                      func.date_format(CaseList.case_stime, '%Y-%m-%d').label('date')).filter(
                    *conditions).group_by('date').all()
            with DBContext('r') as session:
                conditions = []
                conditions.append(CaseList.case_type == "故障")
                conditions.append(CaseList.case_stime >= tostart)
                conditions.append(CaseList.case_etime <= toend)
                lime3 = session.query(func.count(CaseList.case_type),
                                      func.date_format(CaseList.case_stime, '%Y-%m-%d').label('date')).filter(
                    *conditions).group_by('date').all()
            with DBContext('r') as session:
                conditions = []
                conditions.append(or_(CaseList.case_priority == "特急", CaseList.case_ltime >= 240))
                conditions.append(CaseList.case_stime >= tostart)
                conditions.append(CaseList.case_etime <= toend)
                lime4 = session.query(func.count(CaseList.case_priority),
                                      func.date_format(CaseList.case_stime, '%Y-%m-%d').label('date')).filter(
                    *conditions).group_by('date').all()

            # ins_log.read_log('info', lime)
            # ins_log.read_log('info', lime2)
            # 生成开始到结束的日期列表
            begin_date = str(tostart).replace('-', "")
            end_date = str(toend).replace('-', "")
            date_list = []
            begin_date = datetime.datetime.strptime(begin_date, "%Y%m%d")
            end_date = datetime.datetime.strptime(end_date, "%Y%m%d")
            while begin_date <= end_date:
                date_str = begin_date.strftime("%Y-%m-%d")
                date_list.append(date_str)
                begin_date += datetime.timedelta(days=1)
            ins_log.read_log('info', date_list)

            x_lime = date_list
            y_lime = getlist(date_list, lime)
            y_lime2 = getlist(date_list, lime2)
            y_lime3 = getlist(date_list, lime3)
            y_lime4 = getlist(date_list, lime4)
            ins_log.read_log('info', y_lime)
            ins_log.read_log('info', x_lime)
            ins_log.read_log('info', y_lime2)
            # 解决中文显示问题
            from matplotlib.font_manager import FontProperties
            font = FontProperties(fname=r"/usr/share/fonts/china.ttf", size=12)

            plt.figure(figsize=(7, 11))
            plt.xlabel("日期", fontproperties=font)
            plt.ylabel("数量", fontproperties=font)
            plt.plot(x_lime, y_lime, color='r', label="来文", )
            plt.plot(x_lime, y_lime2, color='cyan', label="升级", )
            plt.plot(x_lime, y_lime3, color='b', label="故障", )
            plt.plot(x_lime, y_lime4, color='lime', label="重要工作", )
            plt.title("个案各类情况分布图", fontproperties=font)
            plt.legend(prop=font)  # 添加图例
            plt.xticks(range(0, len(x_lime)), x_lime, rotation=45)

            import time
            time_str = time.strftime('%Y.%m.%d-%H:%M:%S', time.localtime(time.time()))
            Base_DIR = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
            upload_path = '{}/static/report/files/imges/report_form'.format(Base_DIR)
            # filename = "/opt/codo/codo-problem/static/report/files/imges/report_form" + str(time_str)  + ".png"
            filename = upload_path + str(time_str) + ".png"
            plt.savefig(filename)
            # plt.show()
            paragraph = doc.paragraphs[7]  # 获取段落位置
            # p = paragraph.insert_paragraph_before(text="个案情况:每周/每月分布图")
            p = paragraph.insert_paragraph_before(text=" ")
            # p.runs[0].bold = True
            paragraph = doc.paragraphs[8]  # 获取段落位置
            p = paragraph.insert_paragraph_before(text="   ")

            # paragraph = doc.paragraphs[8]  # 获取段落位置
            # paragraph.runs[-1].add_picture(filename, width=shared.Inches(6))  # 在runs的最后一段文字后添加图片/

            # 饼状图
            sum_all = len(data_list)  # 总数
            laiwen = len(laiwen_list) / sum_all  # 来文比例
            upgrade = len(upgrade_list) / sum_all  # 升级比例
            fault = len(fault_list) / sum_all  # 故障比例
            especially = len(especially_list) / sum_all  # 特急比例
            tother = 1 - laiwen - upgrade - fault - especially
            font = FontProperties(fname=r"/usr/share/fonts/china.ttf", size=12)
            label = u'来文', u'升级', u'故障', u'重要工作', u'其他'  # 各类别标签
            color = 'red', 'orange', 'yellow', 'green', 'blue'  # 各类别颜色
            size = [laiwen, upgrade, fault, especially, tother]  # 各类别占比
            explode = (0, 0, 0, 0, 0.2)  # 各类别的偏移半径
            # 绘制饼状图
            plt.figure()  # 清空画板
            plt.figure(figsize=(8, 10))
            patches, l_text, p_text = plt.pie(size, colors=color, explode=explode, labels=label, shadow=True,
                                              autopct='%1.1f%%', startangle=90)
            for t in l_text:
                t.set_fontproperties(font)  # 把每个文本设成中文字体
            # plt.axis('equal')
            plt.title(u'工作完成情况', fontproperties=font, fontsize=12)
            plt.legend(prop=font, loc='lower left', fontsize=12)  # 图例
            # 设置legend的字体大小
            leg = plt.gca().get_legend()
            ltext = leg.get_texts()
            plt.setp(ltext, fontsize=6)

            time_str = time.strftime('%Y.%m.%d-%H:%M:%S', time.localtime(time.time()))
            Base_DIR = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
            upload_path = '{}/static/report/files/imges/report_pie'.format(Base_DIR)
            # filename = "/opt/codo/codo-problem/static/report/files/imges/report_pie" + str(time_str) + ".png"
            filename = upload_path + str(time_str) + ".png"
            plt.savefig(filename)
            # plt.show()

            paragraph = doc.paragraphs[9]  # 获取段落位置
            p = paragraph.insert_paragraph_before(
                text="比例饼图:本月/周完成来文登记" + str(len(laiwen_list)) + "次,平台升级" + str(len(upgrade_list)) + "次，同时处理故障" + str(
                    len(fault_list)) + "条,并汇总处理情况。另，也按领导要求完成其他重要工作" + str(len(especially_list)) + "项。")
            p.runs[0].bold = True
            paragraph = doc.paragraphs[10]  # 获取段落位置
            p = paragraph.insert_paragraph_before(text="  ")
            paragraph = doc.paragraphs[10]  # 获取段落位置
            paragraph.runs[-1].add_picture(filename, width=shared.Inches(6))  # 在runs的最后一段文字后添加图片/
            # 柱状图
            label = ['来文', '升级', '故障', '重要工作', '其他']  # 各类别标签
            plt.figure()  # 清空画板
            plt.figure(figsize=(6, 9))
            sum_other = sum_all - len(laiwen_list) - len(upgrade_list) - len(fault_list) - len(especially_list)
            data = [len(laiwen_list), len(upgrade_list), len(fault_list), len(especially_list), sum_other]
            plt.bar(range(len(data)), data, color=['r', 'g', 'b', 'gold'])  # or `color=['r', 'g', 'b']`

            plt.xticks(range(0, len(label)), label, fontproperties=font, fontsize=12)
            # plt.xlabel("个案情况", fontproperties=font,fontsize=12)
            plt.ylabel("数量", fontproperties=font, fontsize=12)
            plt.title("工作完成情况", fontproperties=font, fontsize=12)
            time_str = time.strftime('%Y.%m.%d-%H:%M:%S', time.localtime(time.time()))
            Base_DIR = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
            upload_path = '{}/static/report/files/imges/report_bar'.format(Base_DIR)
            # filename = "/opt/codo/codo-problem/static/report/files/imges/report_bar" + str(time_str) + ".png"
            filename = upload_path + str(time_str) + ".png"
            plt.savefig(filename)

            paragraph = doc.paragraphs[11]  # 获取段落位置
            p = paragraph.insert_paragraph_before(
                text="个案柱状图:本月/周完成来文登记" + str(len(laiwen_list)) + "次,平台升级" + str(len(upgrade_list)) + "次，同时处理故障" + str(
                    len(fault_list)) + "条,并汇总处理情况。另，也按领导要求完成其他重要工作" + str(len(especially_list)) + "项。")
            p.runs[0].bold = True
            paragraph = doc.paragraphs[12]  # 取段落位置 获
            p = paragraph.insert_paragraph_before(text="  ")
            paragraph = doc.paragraphs[12]  # 获取段落位置
            paragraph.runs[-1].add_picture(filename, width=shared.Inches(5))  # 在runs的最后一段文字后添加图片/
            Base_DIR = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
            upload_path = '{}/static/report/files/'.format(Base_DIR)
            # doc.save(u"/opt/codo/codo-problem/static/report/files/" + tempstr)  # 保存文档
            doc.save(upload_path + tempstr)  # 保存文档
            # http://192.168.2.200:8200/static/report/files/%E7%BB%B4%E6%8A%A4%E7%BB%84%E5%B7%A5%E4%BD%9C%E6%8A%A5%E5%91%8A_admin[20200601-20200630].docx
            # ins_log.read_log('info',self.request.host)
            urlstr = "http://" + self.request.host + "/static/report/files/" + tempstr
            self.write(dict(code=0, msg='获取報告成功', count=1, data=urlstr))
        else:
            self.write(dict(code=-1, msg='没有相关数据', count=0, data=[]))


def getlist(date_list, lime):
    sum_list = []
    sum_flag = 0
    for d in date_list:
        for k in lime:
            if d == str(k[1]):
                g = k
                sum_flag = 1
        if sum_flag == 1:
            sum_list.append(g[0])
        else:
            sum_list.append(0)
        sum_flag = 0
    # ins_log.read_log('info', sum_list)
    return sum_list


class caseDelete(BaseHandler):
    def delete(self, *args, **kwargs):
        data = json.loads(self.request.body.decode("utf-8"))
        toid = int(data.get('id', None))  # id号
        with DBContext('w', None, True) as session:
            session.query(CaseList).filter(CaseList.id == toid).delete(synchronize_session=False)
            session.commit()
        return self.write(dict(code=0, msg='删除成功'))


class getBarHandler(BaseHandler):
    def get(self, *args, **kwargs):
        data_list = []
        nickname = self.get_current_nickname()
        tostart = self.get_argument('startdate', strip=True) + " " + "00:00:00"  # 要查询的关键字
        toend = self.get_argument('enddate', strip=True) + " " + "23:59:59"  # 要查询的关键字

        with DBContext('r') as session:
            conditions = []
            conditions.append(CaseList.case_stime >= tostart)
            conditions.append(CaseList.case_etime <= toend)
            todata = session.query(CaseList).filter(*conditions).order_by(CaseList.ctime.desc()).all()
        for msg in todata:
            case_dict = {}
            data_dict = model_to_dict(msg)
            case_dict["id"] = data_dict["id"]
            case_dict["case_num"] = data_dict["case_num"]
            case_dict["case_obj"] = data_dict["case_obj"]
            case_dict["demand_unit"] = data_dict["demand_unit"]
            case_dict["case_details"] = data_dict["case_details"]
            case_dict["case_type"] = data_dict["case_type"]
            case_dict["case_ltime"] = data_dict["case_ltime"]
            case_dict["case_name"] = data_dict["case_name"]
            case_dict["case_status"] = data_dict["case_status"]
            case_dict["case_priority"] = data_dict["case_priority"]
            case_dict["demander"] = data_dict["demander"]
            case_dict["case_executor"] = data_dict["case_executor"]
            case_dict["case_source"] = data_dict["case_source"]
            case_dict["case_details"] = data_dict["case_details"]
            case_dict["case_stime"] = str(data_dict["case_stime"])
            case_dict["case_etime"] = str(data_dict["case_etime"])
            case_dict["case_creator"] = data_dict["case_creator"]
            data_list.append(case_dict)

        if len(data_list) > 0:
            laiwen_list = []  # 来问
            upgrade_list = []  # 升级
            fault_list = []  # 故障
            plan_list = []  # 计划工作完成情况
            especially_list = []  # 特急/耗时长的情况
            sudden_list = []  # 其他突发工作
            for i in data_list:
                if i["case_source"] == '来文':
                    laiwen_list.append(i)
                if i["case_type"] == '应用升级':
                    upgrade_list.append(i)
                if i["case_type"] == '故障':
                    fault_list.append(i)
                if i["case_priority"] == '特急' or int(i["case_ltime"]) >= 240:
                    especially_list.append(i)
            other_all = len(data_list) - len(laiwen_list) - len(upgrade_list) - len(fault_list) - len(especially_list)
            bar_list = [
                {"来文": len(laiwen_list), "应用升级": len(upgrade_list), "故障": len(fault_list), "重要工作": len(especially_list),
                 "其他": other_all}]
            pie_list = [{"name": "来文", "value": len(laiwen_list)}, {"name": "应用升级", "value": len(upgrade_list)}, \
                        {"name": "故障", "value": len(fault_list)}, {"name": "重要工作", "value": len(especially_list)},
                        {"name": "其他", "value": other_all}],

            self.write(dict(code=0, msg='获取報告成功', count=1, data=bar_list, list=pie_list))
        else:
            self.write(dict(code=-1, msg='没有相关数据', count=0, data=[]))

class getdemanderList(BaseHandler):
    def get(self, *args, **kwargs):
        data_list = []
        tokey = self.get_argument('key', default=None, strip=True)
        value = self.get_argument('value', default=None, strip=True)

        user_list = []
        with DBContext('r') as session:
            conditions = []
            if tokey == "company":
                conditions.append(CaseList.demand_unit == value)
            else:
                conditions.append(CaseList.demand_unit == tokey)
                conditions.append(CaseList.demander == value)
            todata = session.query(CaseList).filter(*conditions).order_by(CaseList.ctime.desc()).all()
            tocount = session.query(CaseList).filter(*conditions).count()

            # todata = session.query(CaseList).filter(CaseList.demander == value).order_by(CaseList.ctime.desc()).all()
            # tocount = session.query(CaseList).filter(CaseList.demander == value).count()

        for msg in todata:
            case_dict = {}
            data_dict = model_to_dict(msg)
            case_dict["id"] = data_dict["id"]
            case_dict["case_num"] = data_dict["case_num"]
            case_dict["case_obj"] = data_dict["case_obj"]
            case_dict["demand_unit"] = data_dict["demand_unit"]
            case_dict["case_details"] = data_dict["case_details"]
            case_dict["case_type"] = data_dict["case_type"]
            case_dict["case_ltime"] = data_dict["case_ltime"]
            case_dict["case_name"] = data_dict["case_name"]
            case_dict["case_status"] = data_dict["case_status"]
            case_dict["case_priority"] = data_dict["case_priority"]
            case_dict["demander"] = data_dict["demander"]
            case_dict["case_executor"] = data_dict["case_executor"]
            case_dict["case_source"] = data_dict["case_source"]
            case_dict["case_details"] = data_dict["case_details"]
            case_dict["case_stime"] = str(data_dict["case_stime"])
            case_dict["case_etime"] = str(data_dict["case_etime"])
            case_dict["case_creator"] = data_dict["case_creator"]
            case_dict["ctime"] = str(data_dict["ctime"])
            data_list.append(case_dict)

        if len(data_list) > 0:
            self.write(dict(code=0, msg='获取成功', count=tocount, data=data_list))
        else:
            self.write(dict(code=-1, msg='没有相关数据', count=0, data=[]))

class gefileturl(BaseHandler):
    def get(self, *args, **kwargs):
        urlstr = "http://" + self.request.host + "/static/report/template/" + "模板文件.xls"
        self.write(dict(code=0, msg='获取報告成功', count=1, data=urlstr))

penson_urls = [
    (r"/v2/case/add/", CaseListHandler),
    (r"/v2/case/list/", getCaseListHandler),
    (r"/v2/case/getfile/", getCasefileHandler),
    (r"/v2/case/delete/", caseDelete),
    (r"/v2/case/getbar/", getBarHandler),
    (r"/v2/case/demanderlist/", getdemanderList),
    (r"/v2/case/geturl/", gefileturl),
]

if __name__ == "__main__":
    pass
